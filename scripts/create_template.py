"""Script to create the mutual NDA .docx template with Jinja2 placeholders.

Run this once to generate nda_app/templates/mutual_nda.docx.
Uses python-docx for proper formatting, then docxtpl renders the placeholders.
"""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from docx.oxml.ns import qn


def set_style_font(style, name="Times New Roman", size=11):
    style.font.name = name
    style.font.size = Pt(size)


def add_heading_text(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = "Times New Roman"
    return h


def add_body(doc, text):
    p = doc.add_paragraph(text)
    p.style = doc.styles["Normal"]
    return p


def add_section_heading(doc, number, title):
    p = doc.add_paragraph()
    run = p.add_run(f"{number}. {title}")
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(11)
    p.space_after = Pt(6)
    return p


def add_signature_block(doc, party_label, name_var, signer_name_var, signer_title_var):
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run(f"{party_label}: {name_var}")
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(11)

    doc.add_paragraph()

    for label, value in [
        ("Signature", "___________________________"),
        ("Name", signer_name_var),
        ("Title", signer_title_var),
        ("Date", "___________________________"),
    ]:
        p = doc.add_paragraph()
        run = p.add_run(f"{label}: {value}")
        run.font.name = "Times New Roman"
        run.font.size = Pt(11)
        p.space_after = Pt(2)


def create_template():
    doc = Document()

    # Set default font
    style = doc.styles["Normal"]
    set_style_font(style, "Times New Roman", 11)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.15

    # --- Title ---
    title = doc.add_heading("MUTUAL NON-DISCLOSURE AGREEMENT", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(16)

    # --- Preamble ---
    add_body(
        doc,
        'This Mutual Non-Disclosure Agreement ("Agreement") is entered into '
        'as of {{ effective_date }} ("Effective Date") by and between:',
    )

    add_body(doc, "{{ disclosing_party_name }}")
    add_body(doc, "{{ disclosing_party_address }}")
    add_body(doc, '("Party A")')

    add_body(doc, "and")

    add_body(doc, "{{ receiving_party_name }}")
    add_body(doc, "{{ receiving_party_address }}")
    add_body(doc, '("Party B")')

    add_body(
        doc,
        '(each a "Party" and collectively the "Parties").',
    )

    # --- Recitals ---
    add_heading_text(doc, "RECITALS", level=2)

    add_body(
        doc,
        "WHEREAS, the Parties wish to explore {{ purpose }} "
        '(the "Purpose"), and in connection therewith, each Party may disclose '
        "to the other certain confidential and proprietary information.",
    )

    add_body(
        doc,
        "NOW, THEREFORE, in consideration of the mutual promises and covenants "
        "contained herein, and for other good and valuable consideration, the "
        "receipt and sufficiency of which are hereby acknowledged, the Parties "
        "agree as follows:",
    )

    # --- Section 1: Definition of Confidential Information ---
    add_section_heading(doc, 1, "DEFINITION OF CONFIDENTIAL INFORMATION")

    add_body(
        doc,
        '"Confidential Information" means any and all non-public, proprietary, '
        "or confidential information disclosed by either Party (the "
        '"Disclosing Party") to the other Party (the "Receiving Party"), '
        "whether orally, in writing, electronically, or by any other means, "
        "and whether or not marked or designated as confidential, including "
        "but not limited to: trade secrets, business plans, financial "
        "information, customer lists, technical data, product designs, "
        "software, inventions, processes, know-how, marketing strategies, "
        "and any other information that reasonably should be understood to be "
        "confidential given the nature of the information and the "
        "circumstances of disclosure.",
    )

    # --- Section 2: Exclusions ---
    add_section_heading(doc, 2, "EXCLUSIONS FROM CONFIDENTIAL INFORMATION")

    add_body(
        doc,
        "Confidential Information does not include information that: "
        "(a) is or becomes publicly available through no fault of or action "
        "by the Receiving Party; (b) was rightfully in the Receiving Party's "
        "possession prior to disclosure by the Disclosing Party, without "
        "restriction on disclosure; (c) is rightfully received by the "
        "Receiving Party from a third party without restriction on disclosure "
        "and without breach of any obligation of confidentiality; or "
        "(d) is independently developed by the Receiving Party without use "
        "of or reference to the Disclosing Party's Confidential Information.",
    )

    # --- Section 3: Obligations ---
    add_section_heading(doc, 3, "OBLIGATIONS OF RECEIVING PARTY")

    add_body(
        doc,
        "The Receiving Party shall: (a) hold the Confidential Information "
        "in strict confidence; (b) not disclose the Confidential Information "
        "to any third party except as permitted under Section 4; (c) use the "
        "Confidential Information solely for the Purpose; and (d) protect "
        "the Confidential Information using at least the same degree of care "
        "it uses to protect its own confidential information, but in no event "
        "less than reasonable care.",
    )

    # --- Section 4: Permitted Disclosures ---
    add_section_heading(doc, 4, "PERMITTED DISCLOSURES")

    add_body(
        doc,
        "The Receiving Party may disclose Confidential Information to its "
        "employees, officers, directors, professional advisors, and affiliates "
        '(collectively, "Representatives") who: (a) have a need to know such '
        "information for the Purpose; and (b) are bound by confidentiality "
        "obligations at least as protective as those contained in this "
        "Agreement. The Receiving Party shall be responsible for any breach "
        "of this Agreement by its Representatives, including its affiliates.",
    )

    # --- Section 5: Compelled Disclosure ---
    add_section_heading(doc, 5, "COMPELLED DISCLOSURE")

    add_body(
        doc,
        "If the Receiving Party is compelled by law, regulation, or legal "
        "process to disclose Confidential Information, the Receiving Party "
        "shall, to the extent legally permitted, provide the Disclosing Party "
        "with prompt written notice of such requirement so that the Disclosing "
        "Party may seek a protective order or other appropriate remedy. The "
        "Receiving Party shall cooperate with the Disclosing Party in seeking "
        "such protective order. If such protective order or other remedy is "
        "not obtained, the Receiving Party shall disclose only that portion "
        "of the Confidential Information that is legally required to be "
        "disclosed.",
    )

    # --- Section 6: Term and Survival ---
    add_section_heading(doc, 6, "TERM AND SURVIVAL")

    add_body(
        doc,
        "This Agreement shall remain in effect for a period of "
        "{{ term_years }} year(s) from the Effective Date "
        '(the "Term"), unless earlier terminated by either Party upon thirty '
        "(30) days' written notice to the other Party. The obligations of "
        "confidentiality under this Agreement shall survive the expiration "
        "or termination of this Agreement for an additional period of "
        "{{ survival_years }} year(s).",
    )

    # --- Section 7: Return or Destruction ---
    add_section_heading(doc, 7, "RETURN OR DESTRUCTION OF CONFIDENTIAL INFORMATION")

    add_body(
        doc,
        "Upon the Disclosing Party's written request or upon expiration or "
        "termination of this Agreement, the Receiving Party shall promptly "
        "return or destroy all Confidential Information and any copies thereof "
        "in its possession or control, and shall provide written certification "
        "of such return or destruction upon request. Notwithstanding the "
        "foregoing, the Receiving Party may retain copies of Confidential "
        "Information to the extent required by applicable law or regulation, "
        "or in accordance with its standard document retention policies for "
        "archival purposes, provided that such retained copies remain subject "
        "to the confidentiality obligations of this Agreement.",
    )

    # --- Section 8: Remedies ---
    add_section_heading(doc, 8, "REMEDIES")

    add_body(
        doc,
        "Each Party acknowledges that a breach of this Agreement may cause "
        "irreparable harm to the Disclosing Party for which monetary damages "
        "would be an inadequate remedy. Accordingly, either Party may seek "
        "injunctive or other equitable relief to enforce this Agreement, "
        "without the necessity of proving actual damages and without waiving "
        "any other rights or remedies available at law or in equity.",
    )

    # --- Section 9: No License or Warranty ---
    add_section_heading(doc, 9, "NO LICENSE OR WARRANTY")

    add_body(
        doc,
        "Nothing in this Agreement grants either Party any license or rights "
        "in the other Party's Confidential Information except as expressly "
        "set forth herein. ALL CONFIDENTIAL INFORMATION IS PROVIDED "
        '"AS IS" WITHOUT WARRANTY OF ANY KIND, EXPRESS OR IMPLIED.',
    )

    # --- Section 10: Governing Law ---
    add_section_heading(doc, 10, "GOVERNING LAW")

    add_body(
        doc,
        "This Agreement shall be governed by and construed in accordance "
        "with the laws of {{ jurisdiction_display_name }}, without regard "
        "to its conflict of laws principles.",
    )

    # --- Section 11: General Provisions ---
    add_section_heading(doc, 11, "GENERAL PROVISIONS")

    add_body(
        doc,
        "(a) Entire Agreement. This Agreement constitutes the entire "
        "agreement between the Parties with respect to the subject matter "
        "hereof and supersedes all prior and contemporaneous agreements, "
        "understandings, negotiations, and discussions, whether oral or "
        "written.",
    )

    add_body(
        doc,
        "(b) Amendment. This Agreement may not be amended or modified "
        "except by a written instrument signed by both Parties.",
    )

    add_body(
        doc,
        "(c) Assignment. Neither Party may assign this Agreement without "
        "the prior written consent of the other Party, except in connection "
        "with a merger, acquisition, or sale of all or substantially all of "
        "its assets.",
    )

    add_body(
        doc,
        "(d) Severability. If any provision of this Agreement is held to be "
        "invalid or unenforceable, the remaining provisions shall continue "
        "in full force and effect.",
    )

    add_body(
        doc,
        "(e) Counterparts. This Agreement may be executed in counterparts, "
        "each of which shall be deemed an original and all of which together "
        "shall constitute one and the same instrument.",
    )

    add_body(
        doc,
        "(f) Notices. All notices under this Agreement shall be in writing "
        "and delivered to the addresses set forth above.",
    )

    # --- Signature Block ---
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run(
        "IN WITNESS WHEREOF, the Parties have executed this Agreement "
        "as of the Effective Date."
    )
    run.font.name = "Times New Roman"
    run.font.size = Pt(11)

    add_signature_block(
        doc,
        "PARTY A",
        "{{ disclosing_party_name }}",
        "{{ disclosing_party_signer_name }}",
        "{{ disclosing_party_signer_title }}",
    )

    add_signature_block(
        doc,
        "PARTY B",
        "{{ receiving_party_name }}",
        "{{ receiving_party_signer_name }}",
        "{{ receiving_party_signer_title }}",
    )

    # Save template
    output_path = Path(__file__).resolve().parent.parent / "nda_app" / "templates" / "mutual_nda.docx"
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    print(f"Template created at: {output_path}")


if __name__ == "__main__":
    create_template()
